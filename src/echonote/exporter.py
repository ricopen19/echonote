"""出力生成モジュール。"""

from __future__ import annotations

from echonote.transcriber import Segment


def segments_to_text(segments: list[Segment]) -> str:
    """セグメントリストをプレーンテキストに変換する。"""
    return "\n".join(s["text"] for s in segments)


def segments_to_transcript(segments: list[Segment]) -> str:
    """タイムスタンプ付きのトランスクリプトテキストを返す。speaker キーがあればラベルを付加する。"""
    lines = []
    for s in segments:
        start = _fmt_time(s["start"])
        end = _fmt_time(s["end"])
        speaker = s.get("speaker", "")
        prefix = f"{speaker}: " if speaker else ""
        lines.append(f"[{start} - {end}] {prefix}{s['text']}")
    return "\n".join(lines)


def to_markdown(content: str) -> str:
    return content


def to_docx(content: str) -> bytes:
    """Markdown テキストを Word (.docx) に変換して bytes で返す。"""
    from docx import Document
    from docx.shared import Pt
    import io

    doc = Document()

    for line in content.splitlines():
        stripped = line.rstrip()
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(stripped[2:], style="List Bullet")
            p.runs[0].font.size = Pt(11)
        elif stripped == "":
            doc.add_paragraph("")
        else:
            doc.add_paragraph(stripped)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _fmt_time(seconds: float) -> str:
    m, s = divmod(int(seconds), 60)
    return f"{m:02d}:{s:02d}"
